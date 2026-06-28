"""
extraction.py — テキスト抽出層（C1）。

PDF / Word(docx) / Markdown / プレーンテキストから正規化テキストを取り出す。
すべてローカルで完結し、外部サービスへ送信しない（RDD §8：インターネットegressなし）。
OCR は対象外（スコープ外）。

抽出テキストは中間生成物であり、呼び出し側はメモリ上で検出パイプラインへ渡し、
ディスクに残さず破棄することを前提とする（D-2：中間生成物の即時削除）。

全角→半角正規化（NFKC）は既定オフ。法務文書の全角数字・全角記号に対する
規則層 recall を上げたい場合に normalize_width=True で有効化する（検出挙動が変わるため明示選択）。
"""

from __future__ import annotations

import io
import os
import unicodedata
from typing import Optional


SUPPORTED_EXT = {".pdf", ".docx", ".md", ".markdown", ".txt", ".text"}


class ExtractionError(Exception):
    pass


class UnsupportedFormat(ExtractionError):
    pass


def detect_kind(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    if ext in (".md", ".markdown"):
        return "markdown"
    if ext in (".txt", ".text"):
        return "text"
    raise UnsupportedFormat(f"未対応の拡張子: {ext}（対応: {sorted(SUPPORTED_EXT)}）")


def _normalize(text: str, normalize_width: bool) -> str:
    # 改行コードを LF に統一
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    if normalize_width:
        # 全角英数記号 → 半角等（NFKC）。検出仕様を変えるため呼び出し側の明示選択時のみ。
        text = unicodedata.normalize("NFKC", text)
    return text


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _extract_docx(data: bytes) -> str:
    import docx
    document = docx.Document(io.BytesIO(data))
    blocks = []
    for para in document.paragraphs:
        blocks.append(para.text)
    # 表中テキストも回収（法務文書は表が多い）
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    blocks.append(cell.text)
    return "\n".join(blocks)


def extract_text(path: str, normalize_width: bool = False) -> str:
    """ファイルパスから正規化テキストを返す。

    対応拡張子のみ。読み込み・抽出はローカル完結。
    """
    kind = detect_kind(path)
    with open(path, "rb") as f:
        data = f.read()

    if kind == "pdf":
        raw = _extract_pdf(data)
    elif kind == "docx":
        raw = _extract_docx(data)
    elif kind in ("markdown", "text"):
        # Markdown はマスキング対象として「生テキスト」を扱う（記法は剥がさない）。
        raw = data.decode("utf-8", errors="replace")
    else:  # 到達しない
        raise UnsupportedFormat(kind)

    return _normalize(raw, normalize_width)


def extract_from_bytes(data: bytes, kind: str, normalize_width: bool = False) -> str:
    """バイト列から抽出（S3オブジェクトをメモリ上で処理する本番経路向け）。

    kind は 'pdf' / 'docx' / 'markdown' / 'text'。
    """
    if kind == "pdf":
        raw = _extract_pdf(data)
    elif kind == "docx":
        raw = _extract_docx(data)
    elif kind in ("markdown", "text"):
        raw = data.decode("utf-8", errors="replace")
    else:
        raise UnsupportedFormat(kind)
    return _normalize(raw, normalize_width)
